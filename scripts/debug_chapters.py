# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import re

doc = Document('./thesis.docx')

# 查看所有段落，找到章节标题
chapter_indices = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name

    # 检测章节标题
    if re.match(r'^第\d+章', text):
        chapter_indices.append(i)
        print(f'{i:3d} | {style:20s} | {text[:80]}')
    elif re.match(r'^\d+\s+\S', text):
        # 可能是章节标题（如"1 绪论"）
        keywords = ['绪论', '数值模拟', '叶轮', '基准', '优化', '结论']
        if any(keyword in text for keyword in keywords):
            chapter_indices.append(i)
            print(f'{i:3d} | {style:20s} | {text[:80]}')

print(f'\n找到 {len(chapter_indices)} 个章节标题')
print(f'章节索引: {chapter_indices}')

# 查看第1章附近的段落
print('\n第1章附近段落（索引100-150）:')
for i in range(100, min(150, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name
    if text and len(text) > 5:
        print(f'{i:3d} | {style:20s} | {text[:100]}')
