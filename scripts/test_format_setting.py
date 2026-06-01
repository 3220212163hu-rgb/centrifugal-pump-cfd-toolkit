# -*- coding: utf-8 -*-
"""
测试格式设置是否正确
"""

from docx import Document
from docx.shared import Pt

def test_format_setting():
    """测试格式设置"""
    # 创建一个新文档
    doc = Document()
    para = doc.add_paragraph()

    # 添加一个run
    run1 = para.add_run("β")
    run1.font.italic = True

    # 添加另一个run
    run2 = para.add_run("2")
    run2.font.subscript = True

    # 添加第三个run
    run3 = para.add_run("决定了叶轮出口环量")

    # 保存文档
    test_path = os.environ.get('TEST_OUTPUT', './test_format.docx')
    doc.save(test_path)

    print(f"测试文档已保存到：{test_path}")

    # 读取文档并验证格式
    doc2 = Document(test_path)
    para2 = doc2.paragraphs[0]

    print("\n验证格式：")
    for i, run in enumerate(para2.runs):
        print(f"Run {i}: '{run.text}'")
        print(f"  斜体: {run.font.italic}")
        print(f"  上角标: {run.font.superscript}")
        print(f"  下角标: {run.font.subscript}")

if __name__ == '__main__':
    test_format_setting()
