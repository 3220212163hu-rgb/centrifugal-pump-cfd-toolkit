# -*- coding: utf-8 -*-
"""
将改写后的文本回填到Word文档中
保留原有格式，特别是：
1. 物理量的斜体格式
2. 参考文献引用的上角标格式
"""

import re
import json
import shutil
from docx import Document
from docx.shared import Pt
from pathlib import Path

def load_placeholder_map(map_file):
    """加载占位符映射表"""
    with open(map_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def restore_placeholders(text, placeholder_map):
    """将占位符还原为原始内容"""
    for placeholder, original in placeholder_map.items():
        text = text.replace(placeholder, original)
    return text

def find_chapter_boundaries(doc, chapter_num):
    """找到章节的起止段落索引"""
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # 检测章节标题（Heading 1 样式）
        if style == 'Heading 1':
            # 检查是否是目标章节
            if re.match(rf'^{chapter_num}\s+', text) or re.match(rf'^第{chapter_num}章', text):
                chapter_start = i
            elif chapter_start is not None:
                # 找到下一章，停止
                chapter_end = i
                break

    if chapter_start is None:
        return None, None

    if chapter_end is None:
        # 如果是最后一章，取到文档末尾
        chapter_end = len(doc.paragraphs)

    return chapter_start, chapter_end

def is_modifiable_paragraph(para):
    """判断段落是否可以修改"""
    text = para.text.strip()
    style = para.style.name

    if not text:
        return False

    # 跳过子节标题
    if style.startswith('Heading'):
        return False

    # 跳过图表题注
    if re.match(r'^(图|表)\s*\d+-\d+', text):
        return False

    # 跳过公式
    if style == '公式' or 'Equation' in style:
        return False

    # 跳过参考文献列表
    if re.match(r'^\[\d+\]', text):
        return False

    # 跳过目录
    if style.startswith('toc'):
        return False

    return True

def is_reference_citation(text):
    """判断是否是参考文献引用，如 [1]、[2,3]、[4-6]"""
    return bool(re.match(r'^\[\d+(?:[,，]\d+)*(?:-\d+)?\]$', text.strip()))

def is_physics_quantity(text):
    """判断是否是物理量（希腊字母或特定符号）"""
    # 希腊字母
    greek_pattern = r'^[αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]$'
    # 特定物理量符号（单个字母，可能是变量）
    var_pattern = r'^[A-Za-z]$'

    text = text.strip()
    if re.match(greek_pattern, text):
        return True
    # 只有单独出现的字母才可能是物理量变量
    if re.match(var_pattern, text) and len(text) == 1:
        # 排除一些不是物理量的单词
        if text.lower() not in ['a', 'the', 'in', 'on', 'at', 'to', 'for', 'of', 'with']:
            return True
    return False

def is_subscript(text):
    """判断是否是下标"""
    # 常见下标模式：数字、字母
    return bool(re.match(r'^[0-9a-zA-Z]+$', text.strip())) and len(text.strip()) <= 3

def set_run_format(run, is_italic=False, is_superscript=False, is_subscript=False):
    """设置run的格式"""
    run.font.italic = is_italic
    run.font.superscript = is_superscript
    run.font.subscript = is_subscript

def rewrite_paragraph_with_format(para, new_text):
    """重写段落，保留格式"""
    # 清空所有run
    for run in para.runs:
        run.text = ""

    # 如果没有run，添加一个
    if not para.runs:
        run = para.add_run(new_text)
        return

    # 获取第一个run作为模板
    first_run = para.runs[0]
    template_font_name = first_run.font.name
    template_font_size = first_run.font.size
    template_bold = first_run.font.bold

    # 将新文本分割为多个部分，识别需要特殊格式的部分
    # 使用正则表达式匹配参考文献引用和物理量
    parts = []
    remaining = new_text

    while remaining:
        # 匹配参考文献引用 [数字] 或 [数字,数字] 或 [数字-数字]
        ref_match = re.search(r'\[\d+(?:[,，]\d+)*(?:-\d+)?\]', remaining)
        # 匹配希腊字母
        greek_match = re.search(r'[αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]', remaining)
        # 匹配可能的物理量变量（单独的字母，后面跟着下标或数字）
        var_match = re.search(r'(?<![a-zA-Z])([A-Za-z])(?=[₀₁₂₃₄₅₆₇₈₉0-9²³⁴⁵⁶⁷⁸⁹⁺⁻=≈])', remaining)

        # 找到最早匹配的位置
        matches = []
        if ref_match:
            matches.append(('ref', ref_match.start(), ref_match.end(), ref_match.group()))
        if greek_match:
            matches.append(('greek', greek_match.start(), greek_match.end(), greek_match.group()))
        if var_match:
            matches.append(('var', var_match.start(), var_match.end(), var_match.group()))

        if not matches:
            # 没有特殊格式，直接添加剩余文本
            if remaining:
                parts.append(('normal', remaining))
            break

        # 按位置排序
        matches.sort(key=lambda x: x[1])

        # 处理第一个匹配
        match_type, start, end, matched_text = matches[0]

        # 添加匹配前的普通文本
        if start > 0:
            parts.append(('normal', remaining[:start]))

        # 添加匹配的文本
        parts.append((match_type, matched_text))

        # 更新剩余文本
        remaining = remaining[end:]

    # 创建runs
    if not parts:
        parts.append(('normal', new_text))

    # 清空所有run
    for run in para.runs:
        run.text = ""

    # 设置第一个run的文本和格式
    first_run = para.runs[0]
    first_run.text = parts[0][1]

    if parts[0][0] == 'ref':
        set_run_format(first_run, is_superscript=True)
    elif parts[0][0] in ['greek', 'var']:
        set_run_format(first_run, is_italic=True)
    else:
        set_run_format(first_run)

    # 设置后续run的文本和格式
    for i, (part_type, part_text) in enumerate(parts[1:], 1):
        if i < len(para.runs):
            run = para.runs[i]
            run.text = part_text

            if part_type == 'ref':
                set_run_format(run, is_superscript=True)
            elif part_type in ['greek', 'var']:
                set_run_format(run, is_italic=True)
            else:
                set_run_format(run)

            # 复制字体格式
            if template_font_name:
                run.font.name = template_font_name
            if template_font_size:
                run.font.size = template_font_size
            if template_bold:
                run.font.bold = template_bold

def backfill_chapter(doc_path, chapter_num, rewritten_text_file, placeholder_map_file, output_path):
    """将改写后的文本回填到指定章节"""
    # 加载文档
    doc = Document(doc_path)

    # 加载占位符映射表
    placeholder_map = load_placeholder_map(placeholder_map_file)

    # 读取改写后的文本
    with open(rewritten_text_file, 'r', encoding='utf-8') as f:
        rewritten_text = f.read()

    # 将改写后的文本按段落分割
    rewritten_paragraphs = rewritten_text.split('\n\n')
    rewritten_paragraphs = [p.strip() for p in rewritten_paragraphs if p.strip()]

    # 找到章节边界
    chapter_start, chapter_end = find_chapter_boundaries(doc, chapter_num)
    if chapter_start is None:
        print(f"未找到第{chapter_num}章")
        return

    print(f"第{chapter_num}章：段落索引 {chapter_start} 到 {chapter_end}")

    # 收集可修改的段落
    modifiable_indices = []
    for i in range(chapter_start + 1, chapter_end):
        if is_modifiable_paragraph(doc.paragraphs[i]):
            modifiable_indices.append(i)

    print(f"可修改段落数：{len(modifiable_indices)}")
    print(f"改写后段落数：{len(rewritten_paragraphs)}")

    # 检查段落数量是否匹配
    if len(modifiable_indices) != len(rewritten_paragraphs):
        print(f"警告：段落数量不匹配！")
        print(f"  原始可修改段落：{len(modifiable_indices)}")
        print(f"  改写后段落：{len(rewritten_paragraphs)}")
        # 取较小的数量
        min_count = min(len(modifiable_indices), len(rewritten_paragraphs))
        modifiable_indices = modifiable_indices[:min_count]
        rewritten_paragraphs = rewritten_paragraphs[:min_count]

    # 回填改写后的文本
    for idx, rewritten_para in zip(modifiable_indices, rewritten_paragraphs):
        # 还原占位符
        restored_text = restore_placeholders(rewritten_para, placeholder_map)

        # 重写段落，保留格式
        rewrite_paragraph_with_format(doc.paragraphs[idx], restored_text)

    # 保存文档
    doc.save(output_path)
    print(f"已保存到：{output_path}")

def main():
    # 配置
    thesis_path = Path('./thesis.docx')
    output_dir = Path(os.environ.get('OUTPUT_DIR', './extracted'))
    output_path = Path('./thesis_rewrite.docx')

    # 需要回填的章节
    chapters_to_backfill = [1, 5, 6]

    # 复制原始文档作为工作版本
    shutil.copy2(thesis_path, output_path)
    print(f"已创建工作版本：{output_path}")

    # 逐章回填
    for chapter_num in chapters_to_backfill:
        print(f"\n回填第{chapter_num}章...")

        rewritten_text_file = output_dir / f'chapter_{chapter_num}_rewritten.txt'
        placeholder_map_file = output_dir / 'placeholder_map.json'

        if not rewritten_text_file.exists():
            print(f"  跳过：改写文件不存在 {rewritten_text_file}")
            continue

        backfill_chapter(
            output_path,
            chapter_num,
            rewritten_text_file,
            placeholder_map_file,
            output_path  # 直接覆盖工作版本
        )

    print("\n回填完成！")
    print(f"最终文件：{output_path}")

if __name__ == '__main__':
    main()
