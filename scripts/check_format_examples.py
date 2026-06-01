# -*- coding: utf-8 -*-
"""
查看具体的格式例子，确认格式正确保留
"""

import re
from docx import Document

def main():
    # 读取文档
    doc = Document('./thesis_rewrite.docx')

    # 找到第1章
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if style == 'Heading 1':
            if re.match(r'^1\s+', text) or re.match(r'^第1章', text):
                chapter_start = i
            elif chapter_start is not None:
                chapter_end = i
                break

    if chapter_start is None:
        print("未找到第1章")
        return

    if chapter_end is None:
        chapter_end = len(doc.paragraphs)

    print(f"第1章：段落索引 {chapter_start} 到 {chapter_end}")
    print("\n查看格式例子...")

    # 查看前5个段落的详细格式
    count = 0
    for i in range(chapter_start + 1, chapter_end):
        if count >= 5:
            break

        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name

        if not text or style.startswith('Heading'):
            continue

        print(f"\n段落 {i}: {text[:100]}...")
        print(f"  样式: {style}")

        # 显示每个run的格式
        for j, run in enumerate(para.runs):
            run_text = run.text
            if not run_text.strip():
                continue

            is_italic = run.font.italic
            is_superscript = run.font.superscript
            is_subscript = run.font.subscript

            # 只显示有特殊格式的run
            if is_italic or is_superscript or is_subscript:
                print(f"  Run {j}: '{run_text}'")
                print(f"    斜体: {is_italic}")
                print(f"    上角标: {is_superscript}")
                print(f"    下标: {is_subscript}")

        count += 1

    print("\n查看完成！")

if __name__ == '__main__':
    main()
