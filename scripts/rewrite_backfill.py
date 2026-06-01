# -*- coding: utf-8 -*-
"""
将改写后的文本回填到Word文档中
保持原有格式、公式、图表、题注、交叉引用、目录、参考文献、页眉页脚不变
"""

import re
import json
from docx import Document
from pathlib import Path

def load_placeholder_map(map_file):
    """加载占位符映射表"""
    with open(map_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def restore_placeholders(text, placeholder_map):
    """将占位符还原为原始内容"""
    for placeholder, original in placeholder_map.items():
        text = text.replace(placeholder, original)
    return text

def find_chapter_boundaries(doc, chapter_num):
    """找到章节的起止段落索引"""
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # 检测章节标题（Heading 1 样式）
        if style == 'Heading 1':
            # 检查是否是目标章节
            if re.match(rf'^{chapter_num}\s+', text) or re.match(rf'^第{chapter_num}章', text):
                chapter_start = i
            elif chapter_start is not None:
                # 找到下一章，停止
                chapter_end = i
                break

    if chapter_start is None:
        return None, None

    if chapter_end is None:
        # 如果是最后一章，取到文档末尾
        chapter_end = len(doc.paragraphs)

    return chapter_start, chapter_end

def is_modifiable_paragraph(para):
    """判断段落是否可以修改"""
    text = para.text.strip()
    style = para.style.name

    if not text:
        return False

    # 跳过子节标题
    if style.startswith('Heading'):
        return False

    # 跳过图表题注
    if re.match(r'^(图|表)\s*\d+-\d+', text):
        return False

    # 跳过公式
    if style == '公式' or 'Equation' in style:
        return False

    # 跳过参考文献列表
    if re.match(r'^\[\d+\]', text):
        return False

    # 跳过目录
    if style.startswith('toc'):
        return False

    return True

def backfill_chapter(doc_path, chapter_num, rewritten_text_file, placeholder_map_file, output_path):
    """将改写后的文本回填到指定章节"""
    # 加载文档
    doc = Document(doc_path)

    # 加载占位符映射表
    placeholder_map = load_placeholder_map(placeholder_map_file)

    # 读取改写后的文本
    with open(rewritten_text_file, 'r', encoding='utf-8') as f:
        rewritten_text = f.read()

    # 将改写后的文本按段落分割
    rewritten_paragraphs = rewritten_text.split('\n\n')
    rewritten_paragraphs = [p.strip() for p in rewritten_paragraphs if p.strip()]

    # 找到章节边界
    chapter_start, chapter_end = find_chapter_boundaries(doc, chapter_num)
    if chapter_start is None:
        print(f"未找到第{chapter_num}章")
        return

    print(f"第{chapter_num}章：段落索引 {chapter_start} 到 {chapter_end}")

    # 收集可修改的段落
    modifiable_indices = []
    for i in range(chapter_start + 1, chapter_end):
        if is_modifiable_paragraph(doc.paragraphs[i]):
            modifiable_indices.append(i)

    print(f"可修改段落数：{len(modifiable_indices)}")
    print(f"改写后段落数：{len(rewritten_paragraphs)}")

    # 检查段落数量是否匹配
    if len(modifiable_indices) != len(rewritten_paragraphs):
        print(f"警告：段落数量不匹配！")
        print(f"  原始可修改段落：{len(modifiable_indices)}")
        print(f"  改写后段落：{len(rewritten_paragraphs)}")
        # 取较小的数量
        min_count = min(len(modifiable_indices), len(rewritten_paragraphs))
        modifiable_indices = modifiable_indices[:min_count]
        rewritten_paragraphs = rewritten_paragraphs[:min_count]

    # 回填改写后的文本
    for idx, rewritten_para in zip(modifiable_indices, rewritten_paragraphs):
        # 还原占位符
        restored_text = restore_placeholders(rewritten_para, placeholder_map)

        # 保持原有格式，只替换文本
        para = doc.paragraphs[idx]
        # 保存原有的run格式
        if para.runs:
            # 获取第一个run的格式
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic

            # 清空所有run
            for run in para.runs:
                run.text = ""

            # 设置新的文本
            para.runs[0].text = restored_text

            # 恢复格式
            para.runs[0].font.name = font_name
            para.runs[0].font.size = font_size
            para.runs[0].font.bold = bold
            para.runs[0].font.italic = italic
        else:
            # 如果没有run，直接设置文本
            para.text = restored_text

    # 保存文档
    doc.save(output_path)
    print(f"已保存到：{output_path}")

def main():
    # 配置
    thesis_path = Path('./thesis.docx')
    output_dir = Path('D:/AI/SJK/biyesheji/04-论文降重项目/extracted')
    output_path = Path('./thesis_rewrite.docx')

    # 需要回填的章节
    chapters_to_backfill = [1, 5, 6]

    # 复制原始文档作为工作版本
    import shutil
    shutil.copy2(thesis_path, output_path)
    print(f"已创建工作版本：{output_path}")

    # 逐章回填
    for chapter_num in chapters_to_backfill:
        print(f"\n回填第{chapter_num}章...")

        rewritten_text_file = output_dir / f'chapter_{chapter_num}_rewritten.txt'
        placeholder_map_file = output_dir / 'placeholder_map.json'

        if not rewritten_text_file.exists():
            print(f"  跳过：改写文件不存在 {rewritten_text_file}")
            continue

        backfill_chapter(
            output_path,
            chapter_num,
            rewritten_text_file,
            placeholder_map_file,
            output_path  # 直接覆盖工作版本
        )

    print("\n回填完成！")
    print(f"最终文件：{output_path}")

if __name__ == '__main__':
    main()
