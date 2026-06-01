# -*- coding: utf-8 -*-
"""
验证回填后的Word文档
检查内容是否正确，格式是否保持完整
"""

import re
from docx import Document
from pathlib import Path

def verify_chapter(doc, chapter_num):
    """验证指定章节的内容"""
    # 找到章节标题
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # 检测章节标题（Heading 1 样式）
        if style == 'Heading 1':
            # 检查是否是目标章节
            if re.match(rf'^{chapter_num}\s+', text) or re.match(rf'^第{chapter_num}章', text):
                chapter_start = i
            elif chapter_start is not None:
                # 找到下一章，停止
                chapter_end = i
                break

    if chapter_start is None:
        print(f"  未找到第{chapter_num}章")
        return

    if chapter_end is None:
        # 如果是最后一章，取到文档末尾
        chapter_end = len(doc.paragraphs)

    print(f"  第{chapter_num}章：段落索引 {chapter_start} 到 {chapter_end}")

    # 统计段落
    total_paragraphs = 0
    modifiable_paragraphs = 0
    empty_paragraphs = 0

    for i in range(chapter_start + 1, chapter_end):
        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name

        if not text:
            empty_paragraphs += 1
            continue

        total_paragraphs += 1

        # 检查是否是可修改的段落
        if not style.startswith('Heading') and \
           not re.match(r'^(图|表)\s*\d+-\d+', text) and \
           not style == '公式' and \
           not 'Equation' in style and \
           not re.match(r'^\[\d+\]', text) and \
           not style.startswith('toc'):
            modifiable_paragraphs += 1

    print(f"  总段落数：{total_paragraphs}")
    print(f"  可修改段落数：{modifiable_paragraphs}")
    print(f"  空段落数：{empty_paragraphs}")

def verify_format(doc):
    """验证文档格式是否完整"""
    print("\n格式验证：")

    # 检查是否有MERGEFORMAT
    mergeformat_count = 0
    for para in doc.paragraphs:
        if 'MERGEFORMAT' in para.text:
            mergeformat_count += 1

    if mergeformat_count > 0:
        print(f"  警告：发现 {mergeformat_count} 个MERGEFORMAT")
    else:
        print(f"  [OK] 未发现MERGEFORMAT")

    # 检查是否有丢失的符号
    symbols_to_check = ['β₂', 'Qd', 'Hd', 'D₂', 'b₂', 'η']
    missing_symbols = []

    for para in doc.paragraphs:
        text = para.text
        for symbol in symbols_to_check:
            if symbol in text:
                # 检查是否被错误拆分（如 β 2）
                if len(symbol) > 1 and f'{symbol[0]} {symbol[1]}' in text:
                    missing_symbols.append(symbol)

    if missing_symbols:
        print(f"  警告：发现可能被拆分的符号：{set(missing_symbols)}")
    else:
        print(f"  [OK] 符号格式正确")

def main():
    # 配置
    original_path = Path('./thesis.docx')
    rewritten_path = Path('./thesis_rewrite.docx')

    print("=" * 60)
    print("验证回填后的Word文档")
    print("=" * 60)

    # 验证原始文档
    print("\n原始文档：")
    doc_original = Document(original_path)
    for chapter_num in [1, 5, 6]:
        verify_chapter(doc_original, chapter_num)

    # 验证改写后的文档
    print("\n改写后的文档：")
    doc_rewritten = Document(rewritten_path)
    for chapter_num in [1, 5, 6]:
        verify_chapter(doc_rewritten, chapter_num)

    # 验证格式
    verify_format(doc_rewritten)

    # 检查文件大小
    original_size = original_path.stat().st_size
    rewritten_size = rewritten_path.stat().st_size
    print(f"\n文件大小：")
    print(f"  原始文档：{original_size / 1024:.1f} KB")
    print(f"  改写文档：{rewritten_size / 1024:.1f} KB")
    print(f"  差异：{(rewritten_size - original_size) / 1024:.1f} KB")

    print("\n" + "=" * 60)
    print("验证完成！")
    print("=" * 60)

if __name__ == '__main__':
    main()
