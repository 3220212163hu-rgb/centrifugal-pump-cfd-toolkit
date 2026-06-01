# -*- coding: utf-8 -*-
"""
分析原始文档的格式结构，特别是：
1. 物理量的斜体格式
2. 参考文献引用的上角标格式
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def analyze_paragraph_format(para, para_index):
    """分析段落中每个run的格式"""
    print(f"\n段落 {para_index}: {para.text[:80]}...")
    print(f"  样式: {para.style.name}")

    for i, run in enumerate(para.runs):
        text = run.text
        if not text.strip():
            continue

        # 检查格式
        is_italic = run.font.italic
        is_superscript = run.font.superscript
        is_bold = run.font.bold
        font_name = run.font.name
        font_size = run.font.size

        # 判断是否是物理量或参考文献引用
        is_physics_quantity = bool(re.search(r'[αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]', text))
        is_reference = bool(re.search(r'\[\d+\]', text))

        if is_italic or is_superscript or is_physics_quantity or is_reference:
            print(f"  Run {i}: '{text[:50]}'")
            print(f"    斜体: {is_italic}")
            print(f"    上角标: {is_superscript}")
            print(f"    加粗: {is_bold}")
            print(f"    字体: {font_name}")
            print(f"    字号: {font_size}")
            print(f"    是物理量: {is_physics_quantity}")
            print(f"    是参考文献: {is_reference}")

def main():
    # 读取原始文档
    doc = Document('./thesis.docx')

    # 找到第1章
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if style == 'Heading 1':
            if re.match(r'^1\s+', text) or re.match(r'^第1章', text):
                chapter_start = i
            elif chapter_start is not None:
                chapter_end = i
                break

    if chapter_start is None:
        print("未找到第1章")
        return

    if chapter_end is None:
        chapter_end = len(doc.paragraphs)

    print(f"第1章：段落索引 {chapter_start} 到 {chapter_end}")

    # 分析前5个段落的格式
    count = 0
    for i in range(chapter_start + 1, min(chapter_start + 10, chapter_end)):
        para = doc.paragraphs[i]
        if para.text.strip() and not para.style.name.startswith('Heading'):
            analyze_paragraph_format(para, i)
            count += 1
            if count >= 5:
                break

if __name__ == '__main__':
    main()
