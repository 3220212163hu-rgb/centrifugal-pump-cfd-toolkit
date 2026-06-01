# -*- coding: utf-8 -*-
"""
验证回填后的文档是否正确保留了格式：
1. 参考文献引用是否是上角标
2. 物理量是否是斜体
"""

import re
from docx import Document

def verify_paragraph_format(para, para_index):
    """验证段落的格式"""
    issues = []

    for i, run in enumerate(para.runs):
        text = run.text
        if not text.strip():
            continue

        # 检查参考文献引用
        if re.search(r'\[\d+(?:[,，]\d+)*(?:-\d+)?\]', text):
            if not run.font.superscript:
                issues.append(f"  Run {i}: 参考文献引用 '{text}' 不是上角标")

        # 检查物理量（希腊字母）
        greek_pattern = r'[αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]'
        if re.search(greek_pattern, text):
            if not run.font.italic:
                issues.append(f"  Run {i}: 希腊字母 '{text}' 不是斜体")

        # 检查物理量变量（单独的字母，后面跟着下标或数字）
        var_match = re.search(r'(?<![a-zA-Z])([A-Za-z])(?=[₀₁₂₃₄₅₆₇₈₉0-9²³⁴⁵⁶⁷⁸⁹⁺⁻=≈])', text)
        if var_match:
            letter = var_match.group(1)
            if not run.font.italic:
                issues.append(f"  Run {i}: 物理量变量 '{letter}' 不是斜体")

    return issues

def main():
    # 读取文档
    doc = Document('./thesis_rewrite.docx')

    # 找到第1章
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if style == 'Heading 1':
            if re.match(r'^1\s+', text) or re.match(r'^第1章', text):
                chapter_start = i
            elif chapter_start is not None:
                chapter_end = i
                break

    if chapter_start is None:
        print("未找到第1章")
        return

    if chapter_end is None:
        chapter_end = len(doc.paragraphs)

    print(f"第1章：段落索引 {chapter_start} 到 {chapter_end}")
    print("\n验证格式...")

    # 验证每个段落
    total_issues = 0
    for i in range(chapter_start + 1, chapter_end):
        para = doc.paragraphs[i]
        if para.text.strip() and not para.style.name.startswith('Heading'):
            issues = verify_paragraph_format(para, i)
            if issues:
                print(f"\n段落 {i}: {para.text[:80]}...")
                for issue in issues:
                    print(issue)
                total_issues += len(issues)

    print(f"\n验证完成！发现 {total_issues} 个格式问题。")

if __name__ == '__main__':
    main()
