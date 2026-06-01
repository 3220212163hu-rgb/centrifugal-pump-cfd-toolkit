# -*- coding: utf-8 -*-
"""
将改写后的文本回填到Word文档中
正确保留格式：
1. 物理量变量（β、Q、H、n、D、b、Z等）是斜体
2. 下标（2、d、s等）是下角标
3. 参考文献引用（[1]、[2]等）是上角标
"""

import re
import json
import shutil
from docx import Document
from docx.shared import Pt
from pathlib import Path

def load_placeholder_map(map_file):
    """加载占位符映射表"""
    with open(map_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def restore_placeholders(text, placeholder_map):
    """将占位符还原为原始内容"""
    for placeholder, original in placeholder_map.items():
        text = text.replace(placeholder, original)
    return text

def find_chapter_boundaries(doc, chapter_num):
    """找到章节的起止段落索引"""
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        # 检测章节标题（Heading 1 样式）
        if style == 'Heading 1':
            # 检查是否是目标章节
            if re.match(rf'^{chapter_num}\s+', text) or re.match(rf'^第{chapter_num}章', text):
                chapter_start = i
            elif chapter_start is not None:
                # 找到下一章，停止
                chapter_end = i
                break

    if chapter_start is None:
        return None, None

    if chapter_end is None:
        # 如果是最后一章，取到文档末尾
        chapter_end = len(doc.paragraphs)

    return chapter_start, chapter_end

def is_modifiable_paragraph(para):
    """判断段落是否可以修改"""
    text = para.text.strip()
    style = para.style.name

    if not text:
        return False

    # 跳过子节标题
    if style.startswith('Heading'):
        return False

    # 跳过图表题注
    if re.match(r'^(图|表)\s*\d+-\d+', text):
        return False

    # 跳过公式
    if style == '公式' or 'Equation' in style:
        return False

    # 跳过参考文献列表
    if re.match(r'^\[\d+\]', text):
        return False

    # 跳过目录
    if style.startswith('toc'):
        return False

    return True

def is_english_letter(char):
    """检查字符是否是英文字母"""
    return 'a' <= char <= 'z' or 'A' <= char <= 'Z'

def parse_text_with_format(text):
    """解析文本，识别需要特殊格式的部分

    格式规则（基于原文档分析）：
    1. 希腊字母（β, α, γ, η, ε, ω, μ, ν, ξ, ρ, σ, φ, θ 等）→ 斜体
    2. 单个拉丁字母物理量（Q, H, n, P, R, D, b, Z 等）→ 斜体，但需要满足：
       - 前面是空格、标点或行首
       - 后面是数字下标、空格、标点或行尾
       - 不是单词的一部分（如 CFD, ANSYS, Fluent 等）
       - 不是版本号（如 R2）
    3. 下标数字/字母（β2 中的 2, Qd 中的 d）→ 下角标
    4. 参考文献引用（[1], [2,3], [4-6]）→ 上角标
    """
    parts = []
    i = 0

    while i < len(text):
        # 匹配参考文献引用 [数字] 或 [数字,数字] 或 [数字-数字]
        ref_match = re.match(r'\[\d+(?:[,，]\d+)*(?:-\d+)?\]', text[i:])
        if ref_match:
            parts.append(('ref', ref_match.group()))
            i += ref_match.end()
            continue

        # 匹配希腊字母+下标模式（如 β2、ε、ω）
        greek_match = re.match(r'([αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ])([0-9]+)?', text[i:])
        if greek_match:
            parts.append(('greek', greek_match.group(1)))
            if greek_match.group(2):
                parts.append(('subscript', greek_match.group(2)))
            i += greek_match.end()
            continue

        # 匹配单个拉丁字母物理量+下标模式（如 Qd、Hd、D2、b2、ns）
        # 只有当字母独立出现时才识别为物理量
        # 物理量模式：单个字母 + 可选的数字下标
        var_match = re.match(r'([QHnPRDbZk])([0-9]+)?', text[i:])
        if var_match:
            letter = var_match.group(1)
            subscript_part = var_match.group(2)

            # 检查是否是独立的物理量（不是单词的一部分）
            # 条件1：前面是行首、空格或标点（不是英文字母）
            is_independent_start = (i == 0 or not is_english_letter(text[i-1]))

            # 条件2：后面是行尾、空格、标点或数字下标
            next_pos = i + var_match.end()
            is_independent_end = (next_pos >= len(text) or
                                  not is_english_letter(text[next_pos]) or
                                  subscript_part is not None)

            # 条件3：排除版本号（如 R2、v1.0 等）
            # 版本号通常前面有空格，后面是空格或标点
            is_version = False
            if letter == 'R' and subscript_part:
                # 检查前面是否有软件名称相关词汇
                context_before = text[max(0, i-20):i].lower()
                if any(keyword in context_before for keyword in ['fluent', 'ansys', 'r2', 'r1', 'r3', 'r4', 'r5']):
                    is_version = True

            if is_independent_start and is_independent_end and not is_version:
                parts.append(('var', letter))
                if subscript_part:
                    parts.append(('subscript', subscript_part))
                i += var_match.end()
                continue

        # 匹配物理量+小写字母下标模式（如 Qd、Hd、ns）
        # 只有当字母独立出现时才识别为物理量
        var_with_sub_match = re.match(r'([QHnPRDbZk])([ds])', text[i:])
        if var_with_sub_match:
            letter = var_with_sub_match.group(1)
            subscript_part = var_with_sub_match.group(2)

            # 检查是否是独立的物理量（不是单词的一部分）
            # 条件1：前面是行首、空格或标点（不是英文字母）
            is_independent_start = (i == 0 or not is_english_letter(text[i-1]))

            # 条件2：后面是行尾、空格、标点或字母下标
            next_pos = i + var_with_sub_match.end()
            is_independent_end = (next_pos >= len(text) or
                                  not is_english_letter(text[next_pos]))

            if is_independent_start and is_independent_end:
                parts.append(('var', letter))
                parts.append(('subscript', subscript_part))
                i += var_with_sub_match.end()
                continue

        # 匹配普通字符
        parts.append(('normal', text[i]))
        i += 1

    return parts

def rewrite_paragraph_with_format(para, new_text):
    """重写段落，保留格式"""
    # 获取第一个run作为模板
    if para.runs:
        first_run = para.runs[0]
        template_font_name = first_run.font.name
        template_font_size = first_run.font.size
        template_bold = first_run.font.bold
    else:
        template_font_name = None
        template_font_size = None
        template_bold = None

    # 解析文本，识别需要特殊格式的部分
    parts = parse_text_with_format(new_text)

    # 如果没有解析出任何部分，直接设置文本
    if not parts:
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        return

    # 清空所有run
    for run in para.runs:
        run.text = ""

    # 确保有足够的runs
    while len(para.runs) < len(parts):
        para.add_run("")

    # 设置每个run的文本和格式
    for i, (part_type, part_text) in enumerate(parts):
        run = para.runs[i]
        run.text = part_text

        # 重置格式
        run.font.italic = False
        run.font.superscript = False
        run.font.subscript = False

        # 设置特殊格式
        if part_type == 'ref':
            run.font.superscript = True
        elif part_type in ['greek', 'var']:
            run.font.italic = True
        elif part_type == 'subscript':
            run.font.subscript = True

        # 复制字体格式
        if template_font_name:
            run.font.name = template_font_name
        if template_font_size:
            run.font.size = template_font_size
        if template_bold:
            run.font.bold = template_bold

def backfill_chapter(doc_path, chapter_num, rewritten_text_file, placeholder_map_file, output_path):
    """将改写后的文本回填到指定章节"""
    # 加载文档
    doc = Document(doc_path)

    # 加载占位符映射表
    placeholder_map = load_placeholder_map(placeholder_map_file)

    # 读取改写后的文本
    with open(rewritten_text_file, 'r', encoding='utf-8') as f:
        rewritten_text = f.read()

    # 将改写后的文本按段落分割
    rewritten_paragraphs = rewritten_text.split('\n\n')
    rewritten_paragraphs = [p.strip() for p in rewritten_paragraphs if p.strip()]

    # 找到章节边界
    chapter_start, chapter_end = find_chapter_boundaries(doc, chapter_num)
    if chapter_start is None:
        print(f"未找到第{chapter_num}章")
        return

    print(f"第{chapter_num}章：段落索引 {chapter_start} 到 {chapter_end}")

    # 收集可修改的段落
    modifiable_indices = []
    for i in range(chapter_start + 1, chapter_end):
        if is_modifiable_paragraph(doc.paragraphs[i]):
            modifiable_indices.append(i)

    print(f"可修改段落数：{len(modifiable_indices)}")
    print(f"改写后段落数：{len(rewritten_paragraphs)}")

    # 检查段落数量是否匹配
    if len(modifiable_indices) != len(rewritten_paragraphs):
        print(f"警告：段落数量不匹配！")
        print(f"  原始可修改段落：{len(modifiable_indices)}")
        print(f"  改写后段落：{len(rewritten_paragraphs)}")
        # 取较小的数量
        min_count = min(len(modifiable_indices), len(rewritten_paragraphs))
        modifiable_indices = modifiable_indices[:min_count]
        rewritten_paragraphs = rewritten_paragraphs[:min_count]

    # 回填改写后的文本
    for idx, rewritten_para in zip(modifiable_indices, rewritten_paragraphs):
        # 还原占位符
        restored_text = restore_placeholders(rewritten_para, placeholder_map)

        # 重写段落，保留格式
        rewrite_paragraph_with_format(doc.paragraphs[idx], restored_text)

    # 保存文档
    doc.save(output_path)
    print(f"已保存到：{output_path}")

def main():
    # 配置
    thesis_path = Path('./thesis.docx')
    output_dir = Path('D:/AI/SJK/biyesheji/04-论文降重项目/extracted')
    output_path = Path('./thesis_rewrite.docx')

    # 需要回填的章节
    chapters_to_backfill = [1, 5, 6]

    # 复制原始文档作为工作版本
    shutil.copy2(thesis_path, output_path)
    print(f"已创建工作版本：{output_path}")

    # 逐章回填
    for chapter_num in chapters_to_backfill:
        print(f"\n回填第{chapter_num}章...")

        rewritten_text_file = output_dir / f'chapter_{chapter_num}_rewritten.txt'
        placeholder_map_file = output_dir / 'placeholder_map.json'

        if not rewritten_text_file.exists():
            print(f"  跳过：改写文件不存在 {rewritten_text_file}")
            continue

        backfill_chapter(
            output_path,
            chapter_num,
            rewritten_text_file,
            placeholder_map_file,
            output_path  # 直接覆盖工作版本
        )

    print("\n回填完成！")
    print(f"最终文件：{output_path}")

if __name__ == '__main__':
    main()
