# -*- coding: utf-8 -*-
"""
毕业论文Word生成脚本
功能：按能动学院格式要求生成本科毕业论文Word文档
- 标题样式自动编号
- 公式编辑器（OMML）
- 页眉页脚
- 参考文献自动编号
- 图表自动编号
运行环境：Windows PowerShell, Python 3.14
用法：python generate_thesis.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ========== 配置 ==========
from path_config import PROJECT_DIR
import os
OUTPUT_PATH = os.path.join(PROJECT_DIR, "论文初稿.docx")
# ⚠️ 请根据学校要求修改模板路径
TEMPLATE_PATH = os.environ.get(
    'TEMPLATE_PATH',
    os.path.join(PROJECT_DIR, "格式要求", "模板.docx")
)


# ========== 样式定义 ==========
def setup_styles(doc):
    """配置论文样式"""
    styles = doc.styles

    # ---- 正文样式 ----
    normal = styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)  # 小四号
    normal.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    normal.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    # 设置中文字体
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 标题1样式：章标题 ----
    h1 = styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)  # 三号
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)  # 标题不缩进
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ---- 标题2样式：节标题 ----
    h2 = styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(14)  # 四号
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ---- 标题3样式：小节标题 ----
    h3 = styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(12)  # 小四号
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.first_line_indent = Cm(0)
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    return doc


def setup_page_layout(doc):
    """配置页面布局"""
    section = doc.sections[0]
    # 页边距
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    # 纸张大小
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    return doc


def setup_header_footer(doc):
    """配置页眉页脚"""
    section = doc.sections[0]

    # 页眉：居中显示"本科毕业论文"
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "本科毕业论文"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.style.font.size = Pt(9)  # 小五号
    # 页眉下划线
    pPr = hp._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # 页脚：页码居中
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 插入页码域代码
    run = fp.add_run()
    fldChar1 = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml('<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
    run3._element.append(fldChar2)

    return doc


# ========== 公式插入 ==========
def add_formula(doc, latex_str, caption=""):
    """插入Word公式（OMML格式）

    注意：python-docx不直接支持OMML公式编辑器，
    此函数通过XML方式插入公式域代码。
    复杂公式建议用Mathpix识别后在Word中粘贴。

    简单示例：
    add_formula(doc, "H = (p_2 - p_1)/(rho*g)")
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入公式域代码
    run = p.add_run()
    fldChar_begin = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
    run._element.append(fldChar_begin)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve"> EQ </w:instrText>')
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar_sep = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="separate"/>')
    run3._element.append(fldChar_sep)

    # 公式文本（简化版，复杂公式请用Mathpix）
    run4 = p.add_run(latex_str)
    run4.font.name = 'Cambria Math'
    run4.font.size = Pt(12)

    run5 = p.add_run()
    fldChar_end = parse_xml('<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
    run5._element.append(fldChar_end)

    # 公式编号（右对齐）
    if caption:
        run6 = p.add_run(f"    ({caption})")
        run6.font.size = Pt(12)

    return p


def add_formula_simple(doc, formula_text, number=""):
    """插入简单公式文本（居中，带编号）

    适用于不需要公式编辑器的简单公式展示。
    复杂公式请用Mathpix截图识别后粘贴到Word。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)

    if number:
        run2 = p.add_run(f"    ({number})")
        run2.font.size = Pt(12)

    return p


# ========== 图片插入 ==========
def add_figure(doc, image_path, caption="", width=None):
    """插入图片并添加图题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    if os.path.exists(image_path):
        if width:
            run = p.add_run()
            run.add_picture(image_path, width=width)
        else:
            run = p.add_run()
            run.add_picture(image_path, width=Cm(12))  # 默认宽度12cm
    else:
        run = p.add_run(f"[图片: {image_path}]")
        run.font.color.rgb = RGBColor(255, 0, 0)

    # 图题
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = Cm(0)
        run = cap_p.add_run(caption)
        run.font.size = Pt(10)  # 五号
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return p


# ========== 参考文献条目 ==========
def add_reference(doc, ref_text, number):
    """添加参考文献条目"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)  # 悬挂缩进
    run = p.add_run(f"[{number}] {ref_text}")
    run.font.size = Pt(10)  # 五号
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


# ========== 生成论文 ==========
def generate_thesis():
    """生成论文初稿"""
    doc = Document()

    # 配置
    setup_styles(doc)
    setup_page_layout(doc)
    setup_header_footer(doc)

    # ========== 封面 ==========
    # 标题页
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本科毕业论文")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于数值模拟的离心泵\n参数优化设计与仿真")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(4):
        doc.add_paragraph()

    # 学生信息（请在生成前填写）
    info_items = [
        ("学院：", "（请填写）"),
        ("专业：", "（请填写）"),
        ("姓名：", "（请填写）"),
        ("学号：", "（请填写）"),
        ("指导教师：", "（请填写）"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(label)
        run1.font.size = Pt(14)
        run1.font.name = '宋体'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run2 = p.add_run(value)
        run2.font.size = Pt(14)
        run2.font.name = '宋体'
        run2.underline = True
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 摘要 ==========
    doc.add_heading("摘要", level=1)

    doc.add_paragraph(
        "离心泵作为工业领域应用最广泛的流体机械之一，其运行效率和稳定性直接关系到"
        "整个系统的能耗与安全。进口导叶作为离心泵的重要调节装置，通过改变叶轮进口"
        "预旋条件，对泵的外特性与内部流动结构产生显著影响。本文基于CFD数值模拟方法，"
        "系统研究了不同进口导叶安装角下离心泵的水力性能变化规律，分析了进口预旋对"
        "叶轮内部速度场、压力场及湍流特性的影响机理。"
    )
    doc.add_paragraph(
        "（摘要内容待补充完善，此处为框架示例）"
    )

    # 关键词
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("关键词：")
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2 = p.add_run("离心泵；进口导叶；数值模拟；进口预旋；水力性能")
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 目录页（占位） ==========
    doc.add_heading("目录", level=1)
    doc.add_paragraph("（目录请在Word中自动生成：引用→目录→自动目录）")
    doc.add_page_break()

    # ========== 第1章 绪论 ==========
    doc.add_heading("第1章 绪论", level=1)

    doc.add_heading("1.1 研究背景与意义", level=2)
    doc.add_paragraph(
        "离心泵是石油化工、电力、水利、农业灌溉等领域中最常用的流体输送设备，"
        "其运行效率约占泵系统总能耗的20%以上。随着国家节能减排战略的深入推进，"
        "提高离心泵运行效率、拓宽稳定运行工况范围已成为流体机械领域的重要研究课题。"
        "（待补充）"
    )

    doc.add_heading("1.2 国内外研究现状", level=2)
    doc.add_heading("1.2.1 离心泵进口导叶研究现状", level=3)
    doc.add_paragraph(
        "（此处将基于文献综述内容填充，参考文献[1-15]）"
    )

    doc.add_heading("1.2.2 离心泵内部流动数值模拟研究现状", level=3)
    doc.add_paragraph(
        "（此处将基于文献综述内容填充，参考文献[16-25]）"
    )

    doc.add_heading("1.3 研究内容与方法", level=2)
    doc.add_paragraph(
        "（待补充）"
    )

    # ========== 第2章 离心泵数值模拟理论基础 ==========
    doc.add_heading("第2章 离心泵数值模拟理论基础", level=1)

    doc.add_heading("2.1 计算流体力学基本方程", level=2)
    doc.add_paragraph("流体运动遵循质量守恒、动量守恒和能量守恒三大基本定律。")

    # 连续性方程
    doc.add_paragraph("连续性方程（质量守恒方程）：")
    add_formula_simple(doc, "∂ρ/∂t + ∇·(ρv) = 0", "2-1")

    # 动量方程
    doc.add_paragraph("动量守恒方程（N-S方程）：")
    add_formula_simple(doc, "∂(ρv)/∂t + ∇·(ρvv) = -∇p + ∇·τ + ρg", "2-2")

    doc.add_heading("2.2 湍流模型", level=2)
    doc.add_paragraph(
        "本文采用RNG k-ε湍流模型进行数值计算。该模型在标准k-ε模型的基础上"
        "通过重整化群理论修正了湍流粘度，对旋流和流动分离的预测精度优于标准k-ε模型。"
    )
    add_formula_simple(doc, "∂(ρk)/∂t + ∂(ρkui)/∂xi = ∂/∂xj[(αkμeff)∂k/∂xj] + Gk - ρε", "2-3")

    doc.add_heading("2.3 离心泵基本理论", level=2)
    doc.add_paragraph("离心泵扬程计算公式：")
    add_formula_simple(doc, "H = (p2-p1)/(ρg) + (v2²-v1²)/(2g) + z2-z1", "2-4")

    doc.add_paragraph("离心泵效率计算公式：")
    add_formula_simple(doc, "η = ρgQH / (P_shaft) × 100%", "2-5")

    # ========== 第3章 数值模拟方案 ==========
    doc.add_heading("第3章 数值模拟方案", level=1)

    doc.add_heading("3.1 计算模型", level=2)
    doc.add_paragraph("（待补充：泵参数、几何模型、网格划分）")

    doc.add_heading("3.2 边界条件与求解设置", level=2)
    doc.add_paragraph("（待补充）")

    doc.add_heading("3.3 网格无关性验证", level=2)
    doc.add_paragraph("（待补充）")

    # ========== 第4章 结果与讨论 ==========
    doc.add_heading("第4章 结果与讨论", level=1)

    doc.add_heading("4.1 进口导叶对离心泵外特性的影响", level=2)
    doc.add_paragraph("（基于文献对比归纳分析，待补充）")

    doc.add_heading("4.2 进口预旋对内部流动的影响", level=2)
    doc.add_paragraph("（基于文献对比归纳分析，待补充）")

    doc.add_heading("4.3 不同导叶安装角的性能对比", level=2)
    doc.add_paragraph("（基于文献对比归纳分析，待补充）")

    # ========== 第5章 结论与展望 ==========
    doc.add_heading("第5章 结论与展望", level=1)

    doc.add_heading("5.1 主要结论", level=2)
    doc.add_paragraph("（待补充）")

    doc.add_heading("5.2 研究展望", level=2)
    doc.add_paragraph("（待补充）")

    doc.add_page_break()

    # ========== 参考文献 ==========
    doc.add_heading("参考文献", level=1)

    refs = [
        "徐文斌. 进口导叶对离心泵内部流动特性影响的数值研究[D]. 杭州: 浙江理工大学, 2019.",
        "王宇航. 进口预旋对离心泵内部流动特性影响的数值研究[D]. 杭州: 浙江理工大学, 2022.",
        "曹树良, 谭磊, 朱荣生. 离心泵前置导叶的设计方法及试验分析[C]. 国际农业工程大会, 2010.",
        "（其余参考文献请从正式参考文献表草案.txt中补充）",
    ]

    for i, ref in enumerate(refs, 1):
        add_reference(doc, ref, i)

    doc.add_page_break()

    # ========== 致谢 ==========
    doc.add_heading("致谢", level=1)
    doc.add_paragraph("（待补充）")

    # ========== 保存 ==========
    doc.save(OUTPUT_PATH)
    print(f"论文初稿已生成: {OUTPUT_PATH}")
    print("请在Word中打开，补充内容后保存。")
    print("\n提示：")
    print("1. 目录需在Word中自动生成：引用→目录→自动目录")
    print("2. 公式编号需手动调整或使用域代码")
    print("3. 图片可直接在Word中插入替换占位符")
    print("4. 复杂公式用Mathpix截图识别后粘贴")


if __name__ == "__main__":
    generate_thesis()
