# -*- coding: utf-8 -*-
"""
调试段落格式
"""

import re
from docx import Document

def main():
    # 读取文档
    doc = Document('./thesis_rewrite.docx')

    # 找到第1章
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if style == 'Heading 1':
            if re.match(r'^1\s+', text) or re.match(r'^第1章', text):
                chapter_start = i
            elif chapter_start is not None:
                chapter_end = i
                break

    if chapter_start is None:
        print("未找到第1章")
        return

    if chapter_end is None:
        chapter_end = len(doc.paragraphs)

    # 查看段落 107 的详细信息
    para_index = 107
    para = doc.paragraphs[para_index]
    text = para.text

    print(f"段落 {para_index}: {text[:100]}...")
    print(f"样式: {para.style.name}")
    print(f"\n所有 runs：")

    for i, run in enumerate(para.runs):
        run_text = run.text
        is_italic = run.font.italic
        is_superscript = run.font.superscript
        is_subscript = run.font.subscript

        print(f"Run {i:2d}: '{run_text}'")
        print(f"        斜体: {is_italic}, 上角标: {is_superscript}, 下角标: {is_subscript}")

if __name__ == '__main__':
    main()
