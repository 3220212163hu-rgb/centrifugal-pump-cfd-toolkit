# -*- coding: utf-8 -*-
"""
提取毕业论文正文段落，用占位符保护公式、图表、编号等特殊元素
输出需要改写的纯文本
"""

import os
import re
import json
from docx import Document
from pathlib import Path

# 占位符映射表
placeholder_map = {}
placeholder_counter = {
    'EQ': 0,   # 公式
    'FIG': 0,  # 图
    'TAB': 0,  # 表
    'SYM': 0,  # 符号
    'REF': 0,  # 参考文献
    'NUM': 0,  # 编号
}

def create_placeholder(prefix, original_text):
    """创建占位符并记录映射关系"""
    placeholder_counter[prefix] += 1
    key = f"[{prefix}_{placeholder_counter[prefix]:03d}]"
    placeholder_map[key] = original_text
    return key

def protect_special_elements(text):
    """保护特殊元素，替换为占位符"""
    if not text:
        return text

    # 保护公式编号：如（2-1）、(3-2)、式(4-1)、式（5-2）
    def protect_equation(match):
        return create_placeholder('EQ', match.group(0))
    text = re.sub(r'[（(]\d+-\d+[）)]', protect_equation, text)
    text = re.sub(r'式[（(]\d+-\d+[）)]', protect_equation, text)

    # 保护图编号：如图3-1、图4-2
    def protect_figure(match):
        return create_placeholder('FIG', match.group(0))
    text = re.sub(r'图\s*\d+-\d+', protect_figure, text)

    # 保护表编号：如表3-1、表4-2
    def protect_table(match):
        return create_placeholder('TAB', match.group(0))
    text = re.sub(r'表\s*\d+-\d+', protect_table, text)

    # 保护参考文献引用：如[1]、[2,3]、[4-6]
    def protect_reference(match):
        return create_placeholder('REF', match.group(0))
    text = re.sub(r'\[\d+(?:[,，]\d+)*(?:-\d+)?\]', protect_reference, text)

    # 保护希腊字母和数学符号
    greek_letters = [
        'β₁', 'β₂', 'β1', 'β2', 'α', 'γ', 'δ', 'ε', 'ζ', 'η', 'θ', 'κ',
        'λ', 'μ', 'ν', 'ξ', 'π', 'ρ', 'σ', 'τ', 'φ', 'χ', 'ψ', 'ω',
        'Δ', 'Σ', 'Ω', 'Φ', 'Ψ'
    ]
    for sym in greek_letters:
        if sym in text:
            def protect_sym(match, s=sym):
                return create_placeholder('SYM', s)
            text = text.replace(sym, f'[SYM_{sym}]')
            placeholder_map[f'[SYM_{sym}]'] = sym

    # 保护常用工程符号
    engineering_symbols = ['Qd', 'Hd', 'D₁', 'D₂', 'b₂', 'n', 'η', 'P_s', 'P']
    for sym in engineering_symbols:
        if sym in text:
            placeholder_map[f'[SYM_{sym}]'] = sym
            text = text.replace(sym, f'[SYM_{sym}]')

    # 保护百分比数字（如65.0%）
    def protect_percent(match):
        return create_placeholder('NUM', match.group(0))
    text = re.sub(r'\d+\.?\d*%', protect_percent, text)

    # 保护带单位的数值（如18 m³/h、4 m、900 r/min）
    def protect_value_unit(match):
        return create_placeholder('NUM', match.group(0))
    text = re.sub(r'\d+\.?\d*\s*(m³/h|m³/s|m|mm|rpm|r/min|kW|W|Pa|kPa|MPa|°|℃|K)', protect_value_unit, text)

    return text

def extract_chapter_text(doc_path, chapter_num):
    """提取指定章节的正文段落"""
    doc = Document(doc_path)
    in_chapter = False
    chapter_text = []
    current_paragraph = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测章节标题
        if re.match(rf'^第{chapter_num}章\s', text) or re.match(rf'^{chapter_num}\s+', text):
            in_chapter = True
            continue

        # 检测下一章标题（停止提取）
        if in_chapter and (re.match(r'^第\d+章\s', text) or re.match(r'^\d+\s+', text)):
            if re.match(rf'^第{chapter_num+1}章\s', text) or re.match(rf'^{chapter_num+1}\s+', text):
                break

        # 提取正文段落
        if in_chapter and para.style.name.startswith('Normal') or para.style.name.startswith('正文'):
            # 跳过图表题注
            if re.match(r'^(图|表)\s*\d+-\d+', text):
                continue
            # 跳过公式
            if para.style.name == '公式' or 'Equation' in para.style.name:
                continue
            # 跳过参考文献列表
            if re.match(r'^\[\d+\]', text):
                continue

            chapter_text.append(text)

    return '\n\n'.join(chapter_text)

def extract_all_text(doc_path):
    """提取所有需要改写的正文段落"""
    doc = Document(doc_path)
    all_text = []
    skip_styles = ['公式', 'Equation', 'Heading', 'TOC', '目录', '图题', '表题', '参考文献']

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 跳过特殊样式
        style_name = para.style.name
        if any(skip in style_name for skip in skip_styles):
            continue

        # 跳过图表题注
        if re.match(r'^(图|表)\s*\d+-\d+', text):
            continue

        # 跳过公式
        if re.match(r'^[（(]\d+-\d+[）)]', text):
            continue

        # 跳过参考文献列表
        if re.match(r'^\[\d+\]', text):
            continue

        # 保护特殊元素后添加到列表
        protected_text = protect_special_elements(text)
        all_text.append(protected_text)

    return '\n\n'.join(all_text)

def main():
    # 论文路径
    # ⚠️ 请修改为你的实际论文路径
    thesis_path = Path(os.environ.get('THESIS_PATH', './thesis.docx'))

    # 输出目录
    output_dir = Path(os.environ.get('OUTPUT_DIR', './extracted'))
    output_dir.mkdir(parents=True, exist_ok=True)

    print("正在提取论文内容...")

    # 提取需要重点优化的章节（第1、5、6章）
    chapters_to_extract = [1, 5, 6]

    for chapter_num in chapters_to_extract:
        print(f"\n提取第{chapter_num}章...")
        chapter_text = extract_chapter_text(thesis_path, chapter_num)

        # 保护特殊元素
        protected_text = protect_special_elements(chapter_text)

        # 保存到文件
        output_file = output_dir / f'chapter_{chapter_num}_original.txt'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(protected_text)

        print(f"  已保存到: {output_file}")
        print(f"  字数: {len(protected_text)}")

    # 提取所有正文（用于完整处理）
    print("\n提取所有正文...")
    all_text = extract_all_text(thesis_path)

    # 保存到文件
    output_file = output_dir / 'all_text_original.txt'
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(all_text)

    print(f"  已保存到: {output_file}")
    print(f"  字数: {len(all_text)}")

    # 保存占位符映射表
    map_file = output_dir / 'placeholder_map.json'
    with open(map_file, 'w', encoding='utf-8') as f:
        json.dump(placeholder_map, f, ensure_ascii=False, indent=2)

    print(f"\n占位符映射表已保存到: {map_file}")
    print(f"共创建 {len(placeholder_map)} 个占位符")

if __name__ == '__main__':
    main()
