# -*- coding: utf-8 -*-
"""
详细分析原始文档的格式结构，特别是下角标
"""

import re
from docx import Document

def analyze_paragraph_detailed(para, para_index):
    """详细分析段落中每个run的格式"""
    print(f"\n段落 {para_index}: {para.text[:100]}...")
    print(f"  样式: {para.style.name}")

    for i, run in enumerate(para.runs):
        text = run.text
        if not text.strip():
            continue

        # 检查所有格式属性
        is_italic = run.font.italic
        is_superscript = run.font.superscript
        is_subscript = run.font.subscript
        is_bold = run.font.bold
        font_name = run.font.name
        font_size = run.font.size

        # 只显示有特殊格式的run
        if is_italic or is_superscript or is_subscript:
            print(f"  Run {i}: '{text}'")
            print(f"    斜体: {is_italic}")
            print(f"    上角标: {is_superscript}")
            print(f"    下角标: {is_subscript}")
            print(f"    加粗: {is_bold}")
            print(f"    字体: {font_name}")
            print(f"    字号: {font_size}")

def main():
    # 读取原始文档
    doc = Document('./thesis.docx')

    # 找到第1章
    chapter_start = None
    chapter_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name

        if style == 'Heading 1':
            if re.match(r'^1\s+', text) or re.match(r'^第1章', text):
                chapter_start = i
            elif chapter_start is not None:
                chapter_end = i
                break

    if chapter_start is None:
        print("未找到第1章")
        return

    if chapter_end is None:
        chapter_end = len(doc.paragraphs)

    print(f"第1章：段落索引 {chapter_start} 到 {chapter_end}")

    # 统计格式
    total_runs = 0
    italic_runs = 0
    superscript_runs = 0
    subscript_runs = 0

    # 分析前20个段落
    count = 0
    for i in range(chapter_start + 1, chapter_end):
        if count >= 20:
            break

        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name

        if not text or style.startswith('Heading'):
            continue

        # 统计格式
        for run in para.runs:
            if run.text.strip():
                total_runs += 1
                if run.font.italic:
                    italic_runs += 1
                if run.font.superscript:
                    superscript_runs += 1
                if run.font.subscript:
                    subscript_runs += 1

        # 显示详细的格式信息
        analyze_paragraph_detailed(para, i)
        count += 1

    print(f"\n格式统计：")
    print(f"  总runs数：{total_runs}")
    print(f"  斜体runs：{italic_runs}")
    print(f"  上角标runs：{superscript_runs}")
    print(f"  下角标runs：{subscript_runs}")

if __name__ == '__main__':
    main()
